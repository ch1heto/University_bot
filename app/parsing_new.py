# app/parsing_new.py
"""
Новый парсер документов на базе docx2python.
Заменяет старый parsing.py (2249 строк) и ooxml_lite.py (1521 строка).
Всего ~200 строк вместо 3770!
"""
from __future__ import annotations
import hashlib
from pathlib import Path

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from docx2python import docx2python
import json

# Импортируем утилиты из старого кода
try:
    from .utils import save_bytes, safe_filename, sha256_bytes, ensure_dir
    from .config import Cfg
except:
    # Fallback для тестирования
    def save_bytes(data, filename, dirpath): 
        return str(Path(dirpath) / filename)
    def safe_filename(name, max_len=180): 
        return name
    def sha256_bytes(data): 
        return "hash"
    def ensure_dir(path):
        p = Path(path)
        p.mkdir(parents=True, exist_ok=True)
        return p
    class Cfg:
        UPLOAD_DIR = "./uploads"


# ==================== КОНСТАНТЫ ====================

FIG_CAP_RE = re.compile(
    r"^(?:Рис\.?|Рисунок|Figure|Fig\.?)\s*\.?\s*(?:№\s*)?([A-Za-zА-Яа-я]?\s*\d+(?:[.,]\d+)*)",
    re.IGNORECASE
)

TAB_CAP_RE = re.compile(
    r"^(?:Таблица|Table|Табл\.?)\s*\.?\s*(?:№\s*)?([A-Za-zА-Яа-я]?\s*\d+(?:[.,]\d+)*)",
    re.IGNORECASE
)

HEAD_STYLE_RE = re.compile(r"(heading|заголовок)\s*([1-6])", re.IGNORECASE)


# ==================== ОСНОВНЫЕ ФУНКЦИИ ====================

def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Парсинг .docx файла через python-docx.
    Возвращает словарь с секциями, таблицами, рисунками.
    """
    file_path = str(file_path)
    print(f"DEBUG: Начинаю парсинг файла: {file_path}")

    # 1. ИМПОРТ БИБЛИОТЕКИ
    try:
        from docx import Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
    except ImportError as e:
        print(f"CRITICAL ERROR: Библиотека python-docx не установлена! {e}")
        return {
            'sections': [{'text': 'ОШИБКА: Установите python-docx (pip install python-docx)', 'level': 0}],
            'tables': [],
            'figures': [],
            'metadata': {},
            'raw_text': '',
        }
    
    try:
        doc = Document(file_path)
        
        # Директория для сохранения изображений
        upload_dir = getattr(Cfg, 'UPLOAD_DIR', './uploads')
        images_dir = Path(upload_dir) / 'images'
        images_dir.mkdir(parents=True, exist_ok=True)
        
        # ===================================================================
        # ИЗВЛЕЧЕНИЕ СЕКЦИЙ (ПАРАГРАФЫ)
        # ===================================================================

        def _norm_title_key(s: str) -> str:
            s = re.sub(r"\s+", " ", (s or "").strip().lower())
            s = re.sub(r"[–—−]", "-", s)
            s = re.sub(r"[\"'«»]", "", s)
            s = re.sub(r"[^\wа-яё0-9 -]", "", s)
            return s

        # 0) PRE-PASS: собираем карту подглав из оглавления (строки вида "2.1 ... ..... 15")
        toc_map: dict[str, str] = {}  # normalized_title -> "2.1"
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue
            if _is_toc_line(t):
                m = re.match(r"^(\d+(?:\.\d+)*)\s+(.+?)\.{3,}\s*\d+\s*$", t)
                if m:
                    num = m.group(1).strip()
                    title_toc = m.group(2).strip()
                    if num and title_toc:
                        toc_map[_norm_title_key(title_toc)] = num

        sections = []
        heading_stack: list[str] = []
        current_section_path = "Документ"

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue

            # 1) ВЫКИДЫВАЕМ ОГЛАВЛЕНИЕ ИЗ КОНТЕНТА (оно иначе даёт дубликаты заголовков)
            if _is_toc_line(text):
                continue

            # 2) НЕ ДАЁМ ФОРМУЛАМ/УРАВНЕНИЯМ ПОРТИТЬ СТРУКТУРУ
            # Даже если Word пометил это как Heading — принудительно считаем обычным текстом.
            if _looks_like_formula(text):
                level = 0
                title = ""
                # section_path остаётся текущим
                sections.append({
                    'text': text,
                    'level': 0,
                    'title': "",
                    'element_type': 'text',
                    'section_path': current_section_path,
                })
                continue

            level = 0
            title = ""

            style_name = para.style.name if para.style else ""

            candidate_level = 0
            candidate_title = ""

            # 3) Подтягиваем номер из оглавления, если заголовок в теле без номера/стиля
            # Например в теле: "Организация эмпирического исследования" -> делаем "2.1 Организация..."
            if not re.match(r'^\d+(?:\.\d+)*\b', text):
                key = _norm_title_key(text)
                if key in toc_map:
                    text = f"{toc_map[key]} {text}"

            # 4) Word Heading N — только кандидат
            if style_name.startswith('Heading'):
                try:
                    candidate_level = int(style_name.replace('Heading ', ''))
                except:
                    candidate_level = 1
                candidate_title = text

            # 5) Русские стили "Заголовок N" — тоже кандидат
            elif HEAD_STYLE_RE.search(style_name):
                match = HEAD_STYLE_RE.search(style_name)
                if match:
                    candidate_level = int(match.group(2))
                    candidate_title = text

            # 6) Подтверждение заголовка
            if candidate_level > 0:
                if _is_heading(text):
                    level = candidate_level
                    title = candidate_title
                else:
                    level = 0
                    title = ""
            else:
                if _is_heading(text):
                    level, title = _parse_heading(text)
                else:
                    level = 0
                    title = ""

            # 7) формируем ИЕРАРХИЧЕСКИЙ section_path
            if level > 0:
                try:
                    level = int(level)
                except:
                    level = 1
                if level < 1:
                    level = 1
                if level > 9:
                    level = 9

                heading_stack = heading_stack[: level - 1]
                heading_stack.append(title or text)
                current_section_path = " > ".join(heading_stack)
            else:
                if not current_section_path:
                    current_section_path = "Документ"

            sections.append({
                'text': text,
                'level': level,
                'title': title if title else (text if level > 0 else ""),
                'element_type': 'heading' if level > 0 else 'text',
                'section_path': current_section_path,
            })

            if level > 0:
                print(f"DEBUG HEADING: level={level}, path={current_section_path[:80]}...")


        print(f"DEBUG: Найдено секций (параграфов): {len(sections)}")

        tables = []
        # ===================================================================
        # СНАЧАЛА СОБИРАЕМ ПОДПИСИ ТАБЛИЦ ИЗ ПАРАГРАФОВ
        # ===================================================================
        table_captions_found = []
        for para in doc.paragraphs:
            text = para.text.strip()
            # Ищем "Таблица X.X – Название" или "Таблица X - Название"
            match = re.search(r'[Тт]аблица\s+(\d+(?:\.\d+)?)\s*[-–—:]?\s*(.*)', text)
            if match:
                num = match.group(1)
                title = match.group(2).strip()
                table_captions_found.append({
                    'num': num,
                    'title': title,
                    'full_text': text
                })
        
        print(f"DEBUG: Найдено подписей таблиц: {len(table_captions_found)}")
        
        # ===================================================================
        # ИЗВЛЕЧЕНИЕ ТАБЛИЦ С УМНЫМ СВЯЗЫВАНИЕМ ПОДПИСЕЙ
        # ===================================================================
        # Служебные таблицы (титульный лист, план работы) не имеют подписей
        # Определяем их по содержимому и пропускаем при связывании
        
        caption_index = 0  # Индекс в списке подписей
        
        for table_idx, table in enumerate(doc.tables, start=1):
            rows_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)
            
            if not rows_data:
                continue
            
            # Определяем, является ли таблица служебной (без подписи)
            # Признаки служебной таблицы:
            # 1. Содержит "Исполнитель", "студент", "Научный руководитель"
            # 2. Содержит "№ п/п", "Наименование глав", "Срок выполнения"
            # 3. Первые ячейки пустые (титульная страница)
            
            first_row_text = ' '.join(str(cell) for cell in rows_data[0]).lower()
            all_text = ' '.join(' '.join(str(cell) for cell in row) for row in rows_data[:3]).lower()
            
            is_service_table = (
                'исполнитель' in all_text or
                'научный руководитель' in all_text or
                ('№ п/п' in first_row_text and 'наименование глав' in first_row_text) or
                ('срок' in first_row_text and 'выполнен' in first_row_text) or
                (rows_data[0].count('') > len(rows_data[0]) // 2)  # Больше половины пустых ячеек
            )
            
            if is_service_table:
                # Служебная таблица — не связываем с подписью
                tbl_num = f"служебная_{table_idx}"
                tbl_caption = f"Служебная таблица {table_idx}"
                print(f"DEBUG: Таблица {table_idx} определена как служебная, пропускаем")
            else:
                # Обычная таблица — связываем с подписью
                if caption_index < len(table_captions_found):
                    cap = table_captions_found[caption_index]
                    tbl_num = cap['num']
                    tbl_caption = f"Таблица {cap['num']}"
                    if cap['title']:
                        tbl_caption += f" – {cap['title']}"
                    caption_index += 1
                else:
                    # Fallback: генерируем номер
                    tbl_num = f"доп_{table_idx}"
                    tbl_caption = f'Таблица (дополнительная {table_idx})'
            
            tables.append({
                'id': table_idx,
                'num': tbl_num,
                'caption': tbl_caption,
                'data': rows_data,
                'rows': len(rows_data),
                'cols': len(rows_data[0]) if rows_data else 0,
                'is_service': is_service_table,
            })
            
            print(f"DEBUG: Таблица {tbl_num}: {len(rows_data)} строк, {len(rows_data[0]) if rows_data else 0} колонок")
        
        print(f"DEBUG: Найдено таблиц: {len(tables)} (служебных: {sum(1 for t in tables if t.get('is_service'))})")

        # ===================================================================
        # ИЗВЛЕЧЕНИЕ РИСУНКОВ (ДИАГРАММЫ + ИЗОБРАЖЕНИЯ)
        # ===================================================================
        figures = []
        figure_counter = 0
        
        print(f"DEBUG: Анализ структуры docx (rels count: {len(doc.part.rels)})")

        # ===========================================
        # 1. ИЗВЛЕЧЕНИЕ ДИАГРАММ (CHARTS) - ПРИОРИТЕТ!
        # ===========================================
        try:
            # Импортируем chart_extractor
            try:
                from .chart_extractor import extract_charts_from_docx, get_chart_data_as_text
            except ImportError:
                from chart_extractor import extract_charts_from_docx, get_chart_data_as_text
            
            print(f"DEBUG: Запускаю извлечение диаграмм из {file_path}")
            charts = extract_charts_from_docx(file_path, str(images_dir))
            
            if charts:
                print(f"DEBUG: Найдено {len(charts)} диаграмм!")
                
                for chart in charts:
                    figure_counter += 1
                    
                    # Берём подпись из chart или ищем в секциях
                    caption = chart.get('caption')
                    if not caption or caption.startswith('Рисунок '):
                        found_caption = _find_figure_caption_in_sections(sections, figure_counter)
                        if found_caption:
                            caption = found_caption
                    
                    # Извлекаем номер из подписи
                    fig_num = str(figure_counter)
                    if caption:
                        match = re.search(r'[Рр]ис(?:ун[а-яё]*)?\.?\s*(\d+(?:\.\d+)?)', caption)
                        if match:
                            fig_num = match.group(1)
                    
                    figures.append({
                        'id': figure_counter,
                        'num': fig_num,
                        'caption': caption or f'Рисунок {figure_counter}',
                        'image_path': chart.get('image_path'),
                        'original_name': chart.get('original_name'),
                        'kind': chart.get('kind', 'chart'),
                        'chart_data': chart.get('chart_data'),  # ВАЖНО: данные диаграммы!
                    })
                    
                    print(f"DEBUG: Добавлена диаграмма #{fig_num}: {caption[:50] if caption else 'без подписи'}...")
            else:
                print("DEBUG: Диаграммы не найдены в документе")
                
        except ImportError as e:
            print(f"DEBUG: chart_extractor не найден: {e}")
            print("DEBUG: Убедитесь что файл chart_extractor.py находится в папке app/")
        except Exception as e:
            print(f"DEBUG: Ошибка извлечения диаграмм: {e}")
            import traceback
            traceback.print_exc()

        # ===========================================
        # СНАЧАЛА СОБИРАЕМ ВСЕ ПОДПИСИ РИСУНКОВ ИЗ ПАРАГРАФОВ
        # с позицией параграфа для умного связывания
        # ===========================================
        figure_captions_found = []
        seen_caption_nums = set()  # Для удаления дубликатов
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Ищем "Рисунок X.X" или "Рис. X.X" с возможным описанием
            # Подпись должна НАЧИНАТЬСЯ с "Рис" (не просто упоминание)
            match = re.match(
                r'[Рр]ис(?:ун[а-яё]*)?\.?\s*(\d+(?:\.\d+)?)\s*[-–—:.]?\s*(.*)',
                text
            )
            if match:
                num = match.group(1)
                title = match.group(2).strip()[:100]
                
                # Пропускаем дубликаты — берём первую подпись с непустым названием
                if num not in seen_caption_nums:
                    figure_captions_found.append({
                        'num': num,
                        'title': title,
                        'full_text': text[:200],
                        'para_idx': para_idx,  # Позиция для связывания
                    })
                    seen_caption_nums.add(num)
                elif title:  # Если это дубликат, но с названием — обновляем
                    for cap in figure_captions_found:
                        if cap['num'] == num and not cap['title']:
                            cap['title'] = title
                            cap['full_text'] = text[:200]
                            break
        
        print(f"DEBUG: Найдено подписей рисунков: {len(figure_captions_found)} (уникальных)")
        
        # Создаём словарь подписей по номерам для быстрого поиска
        captions_by_num = {cap['num']: cap for cap in figure_captions_found}
        used_captions = set()  # Какие подписи уже использованы

        # ===========================================
        # 2. ИЗВЛЕЧЕНИЕ ОБЫЧНЫХ ИЗОБРАЖЕНИЙ В ПОРЯДКЕ ДОКУМЕНТА
        # ===========================================
        # ВАЖНО: перебираем изображения в порядке их появления в XML,
        # а не в произвольном порядке из rels (dict)!
        from docx.oxml.ns import qn
        
        raw_images = []
        seen_rel_ids = set()  # Чтобы не добавлять одно изображение дважды
        
        # Перебираем все blip-элементы в порядке документа
        for elem in doc.element.body.iter():
            if elem.tag.endswith('}blip'):
                embed_id = elem.get(qn('r:embed'))
                if embed_id and embed_id not in seen_rel_ids:
                    seen_rel_ids.add(embed_id)
                    
                    try:
                        rel = doc.part.rels.get(embed_id)
                        if not rel:
                            continue
                            
                        reltype = getattr(rel, 'reltype', '') or ''
                        target_ref = getattr(rel, 'target_ref', '') or ''
                        
                        is_image = (
                            'image' in reltype.lower() or
                            'image' in target_ref.lower() or
                            any(target_ref.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.emf', '.wmf'])
                        )
                        
                        if is_image:
                            image_part = rel.target_part
                            image_data = image_part.blob
                            
                            # Фильтруем слишком маленькие изображения (логотипы, иконки)
                            if len(image_data) < 5000:  # < 5KB — скорее всего иконка
                                print(f"DEBUG: Пропускаем маленькое изображение {target_ref} ({len(image_data)} байт)")
                                continue
                            
                            raw_images.append({
                                'rel_id': embed_id,
                                'target_ref': target_ref,
                                'image_data': image_data,
                            })
                            print(f"DEBUG: Найдено изображение #{len(raw_images)}: {target_ref} ({len(image_data)} байт)")
                    except Exception as e:
                        print(f"DEBUG: Ошибка чтения изображения {embed_id}: {e}")
        
        print(f"DEBUG: Найдено изображений для обработки: {len(raw_images)}")
        
        # Теперь связываем изображения с подписями
        # Стратегия: если количество изображений == количеству подписей — связываем по порядку
        # Иначе — пытаемся найти соответствие по номерам из chart_extractor или по порядку с учётом уже использованных
        
        # Сначала добавляем диаграммы (они уже имеют правильные номера)
        chart_nums_used = set()
        for fig in figures:
            if fig.get('num'):
                chart_nums_used.add(fig['num'])
                used_captions.add(fig['num'])
        
        # Определяем какие подписи ещё не использованы
        available_captions = [cap for cap in figure_captions_found if cap['num'] not in used_captions]
        
        print(f"DEBUG: Диаграммы заняли номера: {chart_nums_used}")
        print(f"DEBUG: Доступных подписей для изображений: {len(available_captions)}")
        
        # Связываем изображения с оставшимися подписями
        for img_idx, img in enumerate(raw_images):
            figure_counter += 1
            
            try:
                image_data = img['image_data']
                target_ref = img['target_ref']
                
                ext = os.path.splitext(target_ref)[-1] or ".png"
                img_hash = hashlib.sha256(image_data).hexdigest()[:8]
                safe_name = f"fig_{figure_counter}_{img_hash}{ext}"
                img_path = images_dir / safe_name
                
                with open(img_path, 'wb') as f:
                    f.write(image_data)
                
                # Определяем номер рисунка
                if img_idx < len(available_captions):
                    cap = available_captions[img_idx]
                    fig_num = cap['num']
                    caption = f"Рисунок {cap['num']}"
                    if cap['title']:
                        caption += f" – {cap['title']}"
                    used_captions.add(fig_num)
                    print(f"DEBUG: Изображение #{figure_counter} связано с подписью Рисунок {fig_num}")
                else:
                    # Нет больше подписей — используем порядковый номер
                    fig_num = str(figure_counter)
                    caption = f'Рисунок {figure_counter}'
                    print(f"DEBUG: Изображение #{figure_counter} без подписи (лишнее изображение)")
                
                print(f"DEBUG: Сохранено изображение #{figure_counter} (num={fig_num}): {img_path}")
                
                figures.append({
                    'id': figure_counter,
                    'num': fig_num,
                    'label': fig_num,  # Для совместимости с БД
                    'caption': caption,
                    'image_path': str(img_path.resolve()),
                    'original_name': target_ref,
                    'kind': 'image',
                    'chart_data': None,
                })
                
            except Exception as e:
                print(f"DEBUG: Ошибка извлечения изображения {figure_counter}: {e}")
        
        # ===========================================
        # 3. ПРОВЕРКА НА НЕИЗВЛЕЧЁННЫЕ РИСУНКИ (SmartArt и др.)
        # ===========================================
        # Если подписей больше чем извлечённых рисунков — есть SmartArt или другие объекты
        extracted_nums = {fig['num'] for fig in figures}
        missing_figures = []
        
        for cap in figure_captions_found:
            if cap['num'] not in extracted_nums:
                missing_figures.append(cap)
        
        if missing_figures:
            print(f"DEBUG: ⚠️ ВНИМАНИЕ: {len(missing_figures)} рисунков НЕ извлечены (возможно SmartArt/DrawingML):")
            for mf in missing_figures:
                print(f"DEBUG:   - Рисунок {mf['num']}: {mf['title'][:50] if mf['title'] else '(без названия)'}...")
            print("DEBUG: Эти рисунки могут быть SmartArt-диаграммами, которые не извлекаются как изображения.")
            print("DEBUG: Рекомендация: конвертировать DOCX в PDF для извлечения всех визуальных элементов.")
        
        # Итоговый отчёт
        if not figures:
            print("DEBUG: Рисунков не найдено (ни диаграмм, ни изображений).")
        else:
            charts_count = sum(1 for f in figures if f.get('kind') == 'chart' or f.get('chart_data'))
            images_count = len(figures) - charts_count
            print(f"DEBUG: Всего рисунков: {len(figures)} (диаграмм: {charts_count}, изображений: {images_count})")

        # ===================================================================
        # МЕТАДАННЫЕ И ИТОГ
        # ===================================================================
        core_props = doc.core_properties
        metadata = {
            'author': core_props.author or '',
            'title': core_props.title or '',
            'subject': core_props.subject or '',
            'created': str(core_props.created) if core_props.created else '',
        }
        
        raw_text = '\n'.join([s['text'] for s in sections])
        
        return {
            'sections': sections,
            'tables': tables,
            'figures': figures,
            'metadata': metadata,
            'raw_text': raw_text,
        }
    
    except Exception as e:
        print(f"CRITICAL ERROR в parse_docx: {e}")
        import traceback
        traceback.print_exc()
        
        return {
            'sections': [{'text': f'Ошибка парсинга: {e}', 'level': 0}],
            'tables': [],
            'figures': [],
            'metadata': {},
            'raw_text': '',
        }

def _find_figure_caption_in_sections(sections: list, figure_num: int) -> str:
    """Ищет подпись рисунка в секциях"""
    import re
    
    # Паттерны для поиска подписей
    patterns = [
        rf'[Рр]исунок\s+{figure_num}[\.:\s—–-]+(.+)',
        rf'[Ff]igure\s+{figure_num}[\.:\s—–-]+(.+)',
        rf'[Рр]ис\.\s*{figure_num}[\.:\s—–-]+(.+)',
    ]
    
    for section in sections:
        text = section.get('text', '')
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                caption = match.group(0).strip()
                # Ограничиваем длину подписи (обычно не более 200 символов)
                if len(caption) > 200:
                    caption = caption[:200] + '...'
                return caption
    
    return f'Рисунок {figure_num}'


def parse_doc(file_path: str) -> Dict[str, Any]:
    """
    Парсинг старого формата .doc (через конвертацию в .docx или antiword).
    Fallback: если не можем сконвертировать, возвращаем минимальную структуру.
    """
    file_path = str(file_path)
    
    # Попытка 1: конвертация через LibreOffice (если установлен)
    try:
        docx_path = _convert_doc_to_docx(file_path)
        if docx_path and os.path.exists(docx_path):
            return parse_docx(docx_path)
    except Exception as e:
        print(f"Не удалось сконвертировать .doc: {e}")
    
    # Попытка 2: antiword (если установлен)
    try:
        import subprocess
        text = subprocess.check_output(['antiword', file_path], text=True)
        return {
            'sections': [{'text': text, 'level': 0, 'title': ''}],
            'tables': [],
            'figures': [],
            'metadata': {},
            'raw_text': text,
        }
    except:
        pass
    
    # Fallback: минимальная структура
    return {
        'sections': [{'text': 'Не удалось распарсить .doc файл', 'level': 0, 'title': ''}],
        'tables': [],
        'figures': [],
        'metadata': {},
        'raw_text': '',
    }


def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    Парсинг PDF документа.
    Использует pdfplumber для извлечения текста и таблиц.
    
    Примечание: PDF парсинг сложнее, чем DOCX, потому что PDF - это
    "набор инструкций для рендеринга", а не структурированный документ.
    """
    file_path = str(file_path)
    
    try:
        import pdfplumber
    except ImportError:
        return {
            'sections': [{'text': 'Для парсинга PDF установите: pip install pdfplumber', 'level': 0}],
            'tables': [],
            'figures': [],
            'metadata': {},
            'raw_text': '',
        }
    
    sections = []
    tables = []
    all_text = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            # Метаданные
            metadata = {
                'pages': len(pdf.pages),
                'title': pdf.metadata.get('Title', ''),
                'author': pdf.metadata.get('Author', ''),
            }
            
            for page_num, page in enumerate(pdf.pages, start=1):
                # Извлекаем текст страницы
                text = page.extract_text() or ""
                all_text.append(text)
                
                if text.strip():
                    sections.append({
                        'text': text,
                        'level': 0,
                        'title': '',
                        'page': page_num,
                    })
                
                # Извлекаем таблицы
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables or []):
                    if table:
                        # Очищаем пустые значения
                        cleaned_table = [
                            [cell.strip() if cell else '' for cell in row]
                            for row in table
                        ]
                        
                        tables.append({
                            'id': len(tables) + 1,
                            'num': f"{page_num}.{table_idx + 1}",
                            'caption': f"Таблица на странице {page_num}",
                            'data': cleaned_table,
                            'rows': len(cleaned_table),
                            'cols': len(cleaned_table[0]) if cleaned_table else 0,
                            'page': page_num,
                        })
        
        return {
            'sections': sections,
            'tables': tables,
            'figures': [],  # PDF изображения требуют PyMuPDF (fitz)
            'metadata': metadata,
            'raw_text': '\n\n'.join(all_text),
        }
    
    except Exception as e:
        print(f"Ошибка парсинга PDF: {e}")
        return {
            'sections': [{'text': f'Ошибка парсинга PDF: {e}', 'level': 0}],
            'tables': [],
            'figures': [],
            'metadata': {},
            'raw_text': '',
        }


def save_upload(data: bytes, filename: str, dirpath: str = None) -> str:
    """
    Сохраняет загруженный файл на диск.
    Совместимо со старым API.
    """
    if dirpath is None:
        dirpath = getattr(Cfg, 'UPLOAD_DIR', './uploads')
    
    ensure_dir(dirpath)
    safe_name = safe_filename(filename)
    filepath = Path(dirpath) / safe_name
    
    filepath.write_bytes(data or b"")
    return str(filepath.resolve())


# ==================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ====================

def _safe_get_text(doc_content) -> str:
    """
    Безопасное извлечение текста из документа.
    doc_content.text иногда падает из-за проблем с заголовками/footer.
    """
    try:
        return doc_content.text or ""
    except Exception as e:
        # Fallback: собираем текст из body вручную
        try:
            all_text = []
            for section in doc_content.body:
                for para in section:
                    if isinstance(para, list):
                        text = _flatten_list_to_text(para)
                    else:
                        text = str(para).strip()
                    if text:
                        all_text.append(text)
            return '\n'.join(all_text)
        except:
            return ""

def _extract_sections(doc_content, file_path: str) -> List[Dict[str, Any]]:
    """
    Извлекает секции (параграфы) документа.
    """
    sections = []
    
    # doc_content.body - это вложенная структура:
    # [раздел][параграф/таблица][ячейка (если таблица)][строка]
    for section_idx, section in enumerate(doc_content.body):
        for para_idx, para in enumerate(section):
            # Пропускаем таблицы (они обрабатываются отдельно)
            if isinstance(para, list) and len(para) > 0 and isinstance(para[0], list):
                continue
            
            # Получаем текст параграфа
            if isinstance(para, list):
                text = _flatten_list_to_text(para)
            else:
                text = str(para).strip()
            
            if not text:
                continue
            
            # Определяем уровень заголовка (если есть)
            level = 0
            title = ""
            if _is_heading(text):
                level, title = _parse_heading(text)
            
            sections.append({
                'text': text,
                'level': level,
                'title': title,
                'section_idx': section_idx,
                'para_idx': para_idx,
            })
    
    return sections


def _extract_tables(doc_content) -> List[Dict[str, Any]]:
    """
    Извлекает таблицы из документа.
    """
    tables = []
    table_counter = 0
    
    for section_idx, section in enumerate(doc_content.body):
        for item_idx, item in enumerate(section):
            # Проверяем, является ли элемент таблицей
            if isinstance(item, list) and len(item) > 0 and isinstance(item[0], list):
                # Это таблица!
                table_data = []
                
                for row in item:
                    if isinstance(row, list):
                        # Преобразуем каждую ячейку в строку
                        row_data = [_flatten_list_to_text(cell) if isinstance(cell, list) else str(cell).strip() 
                                    for cell in row]
                        table_data.append(row_data)
                
                # Ищем подпись таблицы (обычно в предыдущем или следующем параграфе)
                caption = _find_table_caption(doc_content.body, section_idx, item_idx)
                
                # Извлекаем номер таблицы из подписи
                table_num = None
                if caption:
                    match = TAB_CAP_RE.match(caption)
                    if match:
                        table_num = match.group(1).strip()
                
                table_counter += 1
                tables.append({
                    'id': table_counter,
                    'num': table_num,
                    'caption': caption,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0,
                    'section_idx': section_idx,
                    'item_idx': item_idx,
                })
    
    return tables


def _extract_figures(doc_content, file_path: str) -> List[Dict[str, Any]]:
    """
    Извлекает рисунки (изображения) из документа.
    """
    figures = []
    figure_counter = 0
    
    # Директория для сохранения изображений
    upload_dir = getattr(Cfg, 'UPLOAD_DIR', './uploads')
    images_dir = ensure_dir(Path(upload_dir) / 'images')
    
    # docx2python автоматически извлекает изображения
    for img_name, img_data in doc_content.images.items():
        figure_counter += 1
        
        # Сохраняем изображение
        img_hash = sha256_bytes(img_data)
        ext = Path(img_name).suffix or '.png'
        safe_name = f"fig_{figure_counter}_{img_hash[:8]}{ext}"
        img_path = images_dir / safe_name
        img_path.write_bytes(img_data)
        
        # Ищем подпись рисунка в тексте документа
        caption = _find_figure_caption(doc_content.text, figure_counter)
        
        # Извлекаем номер рисунка из подписи
        fig_num = None
        if caption:
            match = FIG_CAP_RE.match(caption)
            if match:
                fig_num = match.group(1).strip()
        
        figures.append({
            'id': figure_counter,
            'num': fig_num,
            'caption': caption,
            'image_path': str(img_path.resolve()),
            'original_name': img_name,
            'kind': 'image',  # может быть 'chart', 'diagram', 'photo'
        })
    
    return figures


def _extract_metadata(doc_content) -> Dict[str, Any]:
    """
    Извлекает метаданные документа (автор, дата создания и т.д.)
    """
    metadata = {}
    
    try:
        # docx2python предоставляет properties
        props = getattr(doc_content, 'core_properties', None) or getattr(doc_content, 'properties', None)
        if props:
            metadata = {
                'author': props.get('author', ''),
                'title': props.get('title', ''),
                'subject': props.get('subject', ''),
                'created': props.get('created', ''),
                'modified': props.get('modified', ''),
            }
    except:
        pass
    
    return metadata


# ==================== УТИЛИТЫ ====================

FORMULA_SIGNS_RE = re.compile(r"[=÷±∑√≤≥<>]")
EQ_NUMBER_RE = re.compile(r"\(\s*\d+\s*\)\s*$")
TOC_DOTS_RE = re.compile(r"\.{3,}\s*\d+\s*$")

def _looks_like_formula(text: str) -> bool:
    if FORMULA_SIGNS_RE.search(text):
        return True
    if EQ_NUMBER_RE.search(text) and ("=" in text or "÷" in text):
        return True
    # типичный паттерн "К1 = ..." / "K1 = ..."
    if re.search(r"\b[А-ЯA-Z]\d+\s*=", text):
        return True
    return False

def _is_toc_line(text: str) -> bool:
    # классический лидер точками + номер страницы
    return bool(TOC_DOTS_RE.search(text))

def _flatten_list_to_text(lst) -> str:
    """
    Преобразует вложенный список в плоский текст.
    """
    if isinstance(lst, str):
        return lst.strip()
    elif isinstance(lst, list):
        return ' '.join(_flatten_list_to_text(item) for item in lst).strip()
    else:
        return str(lst).strip()


def _is_heading(text: str) -> bool:
    """
    Проверяет, является ли текст заголовком.
    Версия, которая не превращает формулы и обычные абзацы в заголовки.
    """
    text = (text or "").strip()
    if len(text) < 3 or len(text) > 300:
        return False

    # 0) ЖЁСТКО отсекаем формулы/уравнения (ваш кейс: "К1 = ... (1)")
    # Если есть математические операторы и/или ссылка на формулу "(1)" — это не заголовок.
    if re.search(r'[=÷×+\-*/<>]', text) and re.search(r'\(\s*\d+\s*\)\s*$', text):
        return False
    if re.search(r'[=÷×]', text):
        return False

    # 1) Нумерованные заголовки: "1.2.3 Название" или "1.2.3. Название"
    if re.match(r'^\d+(?:\.\d+)*\.?\s+\S', text):
        # доп. страховка: если это слишком длинная фраза — вероятно абзац/перечень
        return len(text.split()) <= 25

    # 2) "ГЛАВА 1" / "Глава 1" / "CHAPTER 1"
    if re.match(r'^(?:ГЛАВА|CHAPTER)\s*\d+\b', text, re.IGNORECASE):
        return True

    # 3) Полностью заглавные строки — но только если это "текст", а не смесь формульных символов
    # Разрешаем буквы/пробелы/дефисы/кавычки/точки/запятые, без "=" и т.п.
    if text.isupper() and len(text.split()) <= 12 and re.fullmatch(r"[A-ZА-ЯЁ0-9\s«»\"'.,:;()\-]+", text):
        # если заканчивается точкой — скорее абзац
        if not re.search(r'[.]\s*$', text):
            return True

    # 4) Жёсткие ключевые слова разделов (ТОЛЬКО startswith)
    keywords = [
        'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'ВЫВОДЫ', 'ПРИЛОЖЕНИЕ',
        'СПИСОК ЛИТЕРАТУРЫ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
        'БИБЛИОГРАФИЧЕСКИЙ СПИСОК', 'СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ',
        'АННОТАЦИЯ', 'ABSTRACT', 'SUMMARY',
        'INTRODUCTION', 'CONCLUSION', 'REFERENCES', 'APPENDIX'
    ]
    text_upper = text.upper()
    for kw in keywords:
        if text_upper.startswith(kw) and len(text.split()) <= 20:
            return True

    # 5) "1. Название"
    if re.match(r'^\d+\.\s+[А-ЯЁA-Z]', text):
        return True

    return False

def _parse_heading(text: str) -> Tuple[int, str]:
    """
    Определяет уровень заголовка и его текст.
    Улучшенная версия.
    """
    text = text.strip()
    
    # 1) "ГЛАВА 1. Название" или "Глава 1 Название"
    m = re.match(r'^(?:ГЛАВА|Глава|CHAPTER|Chapter)\s*(\d+)[\.\:\s]*(.*)$', text, re.IGNORECASE)
    if m:
        return 1, text  # Главы всегда уровень 1
    
    # 2) Нумерованный заголовок: "1.2.3 Название" или "1.2.3. Название"  
    m = re.match(r'^(\d+(?:\.\d+)*)\.?\s+(.+)$', text)
    if m:
        num_parts = m.group(1).split('.')
        level = len(num_parts)
        return level, text
    
    # 3) Простой номер: "1. Название"
    m = re.match(r'^(\d+)\.\s+(.+)$', text)
    if m:
        return 1, text
    
    # 4) Полностью заглавный текст — уровень 1
    if text.isupper():
        return 1, text
    
    # 5) Ключевые слова — уровень 1
    keywords_l1 = ['ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'ВЫВОДЫ', 'СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ']
    if text.upper() in keywords_l1 or any(text.upper().startswith(k) for k in keywords_l1):
        return 1, text
    
    return 1, text  # По умолчанию уровень 1


def _find_table_caption(body, section_idx: int, item_idx: int) -> Optional[str]:
    """
    Ищет подпись таблицы в соседних параграфах.
    """
    # Проверяем предыдущий параграф
    if item_idx > 0:
        prev_item = body[section_idx][item_idx - 1]
        if isinstance(prev_item, str) or (isinstance(prev_item, list) and not isinstance(prev_item[0], list)):
            text = _flatten_list_to_text(prev_item)
            if TAB_CAP_RE.match(text):
                return text
    
    # Проверяем следующий параграф
    if item_idx < len(body[section_idx]) - 1:
        next_item = body[section_idx][item_idx + 1]
        if isinstance(next_item, str) or (isinstance(next_item, list) and not isinstance(next_item[0], list)):
            text = _flatten_list_to_text(next_item)
            if TAB_CAP_RE.match(text):
                return text
    
    return None


def _find_figure_caption(full_text: str, figure_num: int) -> Optional[str]:
    """
    Ищет подпись рисунка по тексту документа.
    """
    # Ищем строки вида "Рисунок 1 - Название"
    pattern = rf"(?:Рис\.?|Рисунок|Figure)\s*\.?\s*(?:№\s*)?{figure_num}\s*[-–—]\s*(.+)"
    match = re.search(pattern, full_text, re.IGNORECASE)
    if match:
        return match.group(0).strip()
    
    return None


def _convert_doc_to_docx(doc_path: str) -> Optional[str]:
    """
    Конвертирует .doc в .docx через LibreOffice (если установлен).
    """
    import subprocess
    import tempfile
    
    try:
        # Создаём временную директорию
        temp_dir = tempfile.mkdtemp()
        
        # Конвертируем через soffice
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', temp_dir,
            doc_path
        ], check=True, timeout=30)
        
        # Находим сконвертированный файл
        docx_name = Path(doc_path).stem + '.docx'
        docx_path = Path(temp_dir) / docx_name
        
        if docx_path.exists():
            return str(docx_path)
    except:
        pass
    
    return None


# ==================== СОВМЕСТИМОСТЬ СО СТАРЫМ API ====================

def build_index(doc_path: str, output_dir: str = None) -> Dict[str, Any]:
    """
    Создаёт индекс документа (аналог oox_build_index).
    Для совместимости со старым кодом.
    """
    parsed = parse_docx(doc_path)
    
    index = {
        'figures': parsed.get('figures', []),
        'tables': parsed.get('tables', []),
        'sections': parsed.get('sections', []),
        'metadata': parsed.get('metadata', {}),
    }
    
    # Сохраняем индекс в JSON (если нужно)
    if output_dir:
        ensure_dir(output_dir)
        index_path = Path(output_dir) / 'index.json'
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index, f, ensure_ascii=False, indent=2)
    
    return index


def figure_lookup(index: Dict, figure_num: str) -> Optional[Dict]:
    """
    Ищет рисунок по номеру (аналог oox_fig_lookup).
    """
    for fig in index.get('figures', []):
        if fig.get('num') == figure_num:
            return fig
    return None


def table_lookup(index: Dict, table_num: str) -> Optional[Dict]:
    """
    Ищет таблицу по номеру (аналог oox_tbl_lookup).
    """
    for tbl in index.get('tables', []):
        if tbl.get('num') == table_num:
            return tbl
    return None


# ==================== ЭКСПОРТ ====================

__all__ = [
    'parse_docx',
    'parse_doc',
    'parse_pdf',  # ← Добавили
    'save_upload',
    'build_index',
    'figure_lookup',
    'table_lookup',
]