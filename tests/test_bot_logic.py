import re
import pytest

QUESTION = "Что является объектом и предметом исследования?"


def _reset_process_caches():
    # на случай process-level кэша retrieval
    try:
        import app.retrieval as retrieval
        retrieval._DOC_CACHE.clear()
    except Exception:
        pass


def _fake_embeddings(texts):
    # детерминированные вектора, чтобы не дергать сеть
    out = []
    for t in texts:
        h = abs(hash(t)) % 10_000
        vec = [
            (h % 97) / 97.0, (h % 89) / 89.0, (h % 83) / 83.0, (h % 79) / 79.0,
            (h % 73) / 73.0, (h % 71) / 71.0, (h % 67) / 67.0, (h % 61) / 61.0
        ]
        out.append(vec)
    return out


def _extract_obj_subj(intro_text: str):
    obj = None
    subj = None

    # поддерживаем "Объект исследования: ...", "Объектом исследования: ...", тире и т.п.
    m_obj = re.search(r"(?i)\bобъект\w*\s+исследован\w*\s*[:\-–—]?\s*(.+)", intro_text)
    if m_obj:
        obj = m_obj.group(1).strip()

    m_subj = re.search(r"(?i)\bпредмет\w*\s+исследован\w*\s*[:\-–—]?\s*(.+)", intro_text)
    if m_subj:
        subj = m_subj.group(1).strip()

    # вариант "объектом и предметом исследования ..."
    if (obj is None or subj is None):
        m_both = re.search(r"(?i)\bобъект\w*\s+и\s+предмет\w*\s+исследован\w*\s*[:\-–—]?\s*(.+)", intro_text)
        if m_both:
            tail = m_both.group(1).strip()
            obj = obj or tail
            subj = subj or tail

    return obj, subj


def _fake_chat_with_gpt(messages, **kwargs):
    """
    Должен принимать **kwargs (temperature/max_tokens и т.п.).
    """
    ctx = messages[1]["content"]

    if ctx.startswith("[Текст введения]"):
        intro_text = ctx.split("\n", 1)[1] if "\n" in ctx else ""
        obj, subj = _extract_obj_subj(intro_text)

        def q(s: str) -> str:
            # нужна «цитата», иначе sanitizer в answer_semantic_query может обнулить поле
            return f"«{(s or '')[:120]}»" if s else "«…»"

        obj_line = "**Объект:** в тексте ВКР не найдено"
        if obj:
            obj_line = f"**Объект:** {obj} (цитата: {q('Объектом исследования: ' + obj)})"

        subj_line = "**Предмет:** в тексте ВКР не найдено"
        if subj:
            subj_line = f"**Предмет:** {subj} (цитата: {q('Предметом исследования: ' + subj)})"

        # делаем блок достаточно длинным, чтобы не улететь в len<400 fallback
        return "\n".join([
            "## Введение",
            "**Актуальность:** в тексте ВКР не найдено",
            obj_line,
            subj_line,
            "**Цель:** в тексте ВКР не найдено",
            "**Задачи:** в тексте ВКР не найдено",
            "**Гипотеза:** в тексте ВКР не найдено",
        ])

    if ctx.startswith("[Текст главы 1]"):
        return "\n".join([
            "## Глава 1",
            "**Главные идеи:**",
            "- идея 1 (в тексте ВКР не найдено)",
            "- идея 2 (в тексте ВКР не найдено)",
            "- идея 3 (в тексте ВКР не найдено)",
            "**Выводы по главе:** в тексте ВКР не найдено",
        ])

    if ctx.startswith("[Текст главы 2]"):
        return "\n".join([
            "## Глава 2",
            "**Главные идеи:**",
            "- идея 1 (в тексте ВКР не найдено)",
            "- идея 2 (в тексте ВКР не найдено)",
            "- идея 3 (в тексте ВКР не найдено)",
            "**Выводы по главе:** в тексте ВКР не найдено",
        ])

    return "в тексте ВКР не найдено"


def _make_docx(path, intro_obj: str, intro_subj: str):
    """
    Важно: делаем 'ВВЕДЕНИЕ' как обычный абзац (не Heading),
    чтобы воспроизвести ваш кейс, где интро могло "потеряться".
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    p = doc.add_paragraph("ВВЕДЕНИЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True

    doc.add_paragraph(f"Объектом исследования: {intro_obj}.")
    doc.add_paragraph(f"Предметом исследования: {intro_subj}.")

    p1 = doc.add_paragraph("ГЛАВА 1")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p1.runs:
        r.bold = True
    doc.add_paragraph("Текст главы 1 ...")

    p2 = doc.add_paragraph("ГЛАВА 2")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:
        r.bold = True
    doc.add_paragraph("Текст главы 2 ...")

    doc.save(str(path))


@pytest.mark.asyncio
async def test_e2e_intro_role_and_index_order_regression(monkeypatch, tmp_path):
    from app.config import Cfg
    import app.db as db
    import app.parsing as parsing
    import app.retrieval as retrieval
    import app.document_semantic_planner as dsp
    import app.indexing as indexing
    import app.ingest_orchestrator as ingest

    # 1) изоляция БД/кэшей
    Cfg.SQLITE_PATH = str(tmp_path / "test.sqlite")
    _reset_process_caches()

    # 2) моки внешних вызовов
    monkeypatch.setattr(indexing, "embeddings", _fake_embeddings)
    monkeypatch.setattr(dsp, "chat_with_gpt", _fake_chat_with_gpt)

    # 3) готовим два DOCX
    doc_a_path = tmp_path / "A.docx"
    doc_b_path = tmp_path / "B.docx"

    _make_docx(
        doc_a_path,
        intro_obj="предприятие ООО «Лёгкий День»",
        intro_subj="бухгалтерский учет доходов и расходов",
    )
    _make_docx(
        doc_b_path,
        intro_obj="стиль семейного воспитания детей",
        intro_subj="влияние самооценки родителей на выбор стиля семейного воспитания",
    )

    # 4) план по вопросу: хотим, чтобы детектило и obj, и subj (ваш первый баг)
    plan = dsp.build_semantic_plan(QUESTION)
    assert plan.slots.any_slot_requested(), "Слот-детектор не распознал вопрос про объект/предмет."
    assert plan.slots.obj and plan.slots.subj, "Ожидали, что вопрос про объект И предмет выставит оба слота."

    owner_id = 1

    def _index_one(file_path, logical_name: str) -> int:
        doc_id = db.insert_document(owner_id=owner_id, kind="docx", path=str(file_path))
        db.delete_document_chunks(doc_id, owner_id)
        db.delete_document_sections(doc_id)

        sections = parsing.parse_docx(str(file_path))
        sections = ingest.enrich_sections(sections, doc_kind="docx")  # <-- ВАЖНО: тут правильный модуль

        # Проверяем, что intro-role появился хотя бы где-то
        roles = []
        for s in sections:
            attrs = s.get("attrs") if isinstance(s.get("attrs"), dict) else {}
            if attrs.get("role"):
                roles.append(attrs["role"])
        assert "intro" in roles, f"[{logical_name}] После enrich_sections не появился role=intro."

        indexing.index_document(owner_id, doc_id, sections)
        return doc_id

    def _assert_intro_extract(doc_id: int, must_contain: str):
        intro_text = retrieval.get_area_text(owner_id, doc_id, role="intro", max_chars=14000)
        assert intro_text.strip(), "get_area_text(role=intro) вернул пусто."
        assert must_contain in intro_text, f"Во введении не нашли ожидаемую уникальную строку: {must_contain!r}"
        return intro_text

    async def _assert_answer(doc_id: int, must_contain: str, must_not_contain: str):
        ans = await dsp.answer_semantic_query(owner_id, doc_id, QUESTION, plan)
        assert ans, "answer_semantic_query вернул None (значит семантика не сработала или ушла в fallback)."
        assert must_contain in ans, f"В ответе нет ожидаемой строки: {must_contain!r}"
        assert must_not_contain not in ans, "Похоже, ответ “подмешался” из другого документа."

    # ------------------- Сценарий 1: A -> B -> проверяем A и B -------------------
    doc_a = _index_one(doc_a_path, "A")
    _assert_intro_extract(doc_a, "Лёгкий День")
    await _assert_answer(doc_a, "Лёгкий День", "семейного воспитания")

    doc_b = _index_one(doc_b_path, "B")
    _assert_intro_extract(doc_b, "семейного воспитания")
    await _assert_answer(doc_b, "семейного воспитания", "Лёгкий День")

    # повторно проверяем A после индексации B (главная “порядковая” регрессия)
    _assert_intro_extract(doc_a, "Лёгкий День")
    await _assert_answer(doc_a, "Лёгкий День", "семейного воспитания")

    # ------------------- Сценарий 2: B -> A (на новой БД) -------------------
    Cfg.SQLITE_PATH = str(tmp_path / "test2.sqlite")
    _reset_process_caches()

    doc_b2 = _index_one(doc_b_path, "B2")
    _assert_intro_extract(doc_b2, "семейного воспитания")
    await _assert_answer(doc_b2, "семейного воспитания", "Лёгкий День")

    doc_a2 = _index_one(doc_a_path, "A2")
    _assert_intro_extract(doc_a2, "Лёгкий День")
    await _assert_answer(doc_a2, "Лёгкий День", "семейного воспитания")

    # повторно проверяем B2 после индексации A2
    _assert_intro_extract(doc_b2, "семейного воспитания")
    await _assert_answer(doc_b2, "семейного воспитания", "Лёгкий День")
